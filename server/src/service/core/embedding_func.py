from haystack.components.embedders import  SentenceTransformersDocumentEmbedder
import urllib3
from haystack import Pipeline
from haystack.document_stores.in_memory import InMemoryDocumentStore
from haystack.components.converters import HTMLToDocument
from haystack.components.preprocessors import DocumentCleaner
from haystack.components.preprocessors import DocumentSplitter
from haystack.components.writers import DocumentWriter
from haystack_integrations.components.retrievers.qdrant import QdrantEmbeddingRetriever
from haystack.utils import Secret
from haystack_integrations.document_stores.qdrant import QdrantDocumentStore
import os
import pandas as pd
from haystack.dataclasses import Document
from haystack.components.converters import PyPDFToDocument
from haystack.components.converters import TextFileToDocument
import docx  
from haystack.document_stores.types import DuplicatePolicy
from haystack_integrations.components.embedders.cohere import CohereDocumentEmbedder, CohereTextEmbedder

file_path = ".\courses.csv"
url_cloud = "https://f15cf5fc-0771-4b8a-aad5-c4f5c6ae1f1d.us-east4-0.gcp.cloud.qdrant.io:6333"
api_key = "U5tzMbWaGxk3wDvR9yzHCvnFVsTXosi5BR7qFcb7X_j7JOmo4L7RBA"
# index_name = "ThongTinKhoaHoc_Cohere"
# model_name = "embed-multilingual-v3.0"
index_name = "ThongTinKhoaHoc_Cohere"
model_name = "intfloat/multilingual-e5-large-instruct"
#embedding_dim = 768
embedding_dim = 	1024
split_by="word"
split_length=200 
split_overlap = 100

# Call instance qdrant cloud
def load_store(
    index_name: str = index_name, url: str = url_cloud, token: str = api_key, embedding_dim: int = embedding_dim
) -> QdrantDocumentStore:

    return QdrantDocumentStore(
        index=index_name,
        url=url,
        api_key=Secret.from_token(token=token), embedding_dim= embedding_dim
)

def get_name_format_file(filepath: str = file_path):
    file_name = os.path.basename(filepath)
    file = os.path.splitext(file_name)
    return {'file_name':file_name, 'split_name': file}

def embedding_content_fromURL( url: str, index_name:str = index_name):
    # request to url to get data which have html format
    http = urllib3.PoolManager()
    resp = http.request("Get", url)
    # store in file html
    save_to = ".//upload//Temp.html"
    with open(os.path.join(save_to), 'wb') as f:
        f.write(resp.data)

    #Get content from url
    pipeline = Pipeline()
    pipeline.add_component("converter", HTMLToDocument())
    pipeline.add_component("cleaner", DocumentCleaner())
    pipeline.add_component("splitter", DocumentSplitter(split_by=split_by, split_length=split_length, split_overlap = split_overlap))
    #pipeline.add_component("writer", DocumentWriter(document_store=document_store))
    pipeline.connect("converter", "cleaner")
    pipeline.connect("cleaner", "splitter")
    #pipeline.connect("splitter", "writer")
    res = pipeline.run({"converter": {"sources": [save_to]}})

    # split content and store in list documents
    docu = res['splitter']['documents']

    #load qdrant cloud
    document_store = load_store(index_name=index_name)
    # init embedder
    # doc_embedder = CohereDocumentEmbedder(model=model_name)
    doc_embedder = SentenceTransformersDocumentEmbedder( model=model_name)   
    doc_embedder.warm_up()
    ## Use embedder Embedding file document for Fetch và Indexing
    docs_with_embeddings = doc_embedder.run(docu)
    document_store.write_documents(docs_with_embeddings["documents"], policy=DuplicatePolicy.OVERWRITE)
    
    return "Success!"
  


def embedding_txt(filepath: str = "test.txt", index_name: str = index_name):

    document_store = load_store(index_name=index_name)
    pipeline = Pipeline()
    pipeline.add_component("converter", TextFileToDocument())
    pipeline.add_component("cleaner", DocumentCleaner())
    pipeline.add_component("splitter", DocumentSplitter(split_by=split_by, split_length=split_length, split_overlap = split_overlap))
    pipeline.connect("converter", "cleaner")
    pipeline.connect("cleaner", "splitter")

    res = pipeline.run({"converter": {"sources": [filepath]}})
    docu = res['splitter']['documents']
        # init embedder
    # doc_embedder = CohereDocumentEmbedder(model=model_name)
    doc_embedder =SentenceTransformersDocumentEmbedder( model=model_name)  
    doc_embedder.warm_up()
    ## Use embedder Embedding file document for Fetch và Indexing
    docs_with_embeddings = doc_embedder.run(docu)
    document_store.write_documents(docs_with_embeddings["documents"], policy=DuplicatePolicy.OVERWRITE)
    return "Success!"


def embedding_docx(file_path: str, index_name: str = index_name): 
  doc = docx.Document(file_path) 
  file = get_name_format_file(file_path)['split_name']
  # print the list of paragraphs in the document 
  print('List of paragraph objects:->>>') 
  document_text= ""
  for n in doc.paragraphs:
    document_text = f"{document_text} \n {n.text}"

  new_file = os.path.join(os.path.dirname(file_path),f"{file[0]}.txt")
  with open(new_file, "w", encoding="utf-8") as file:
          file.write(document_text)

  embedding_txt(new_file,index_name)
  return "Success!"

    



# Embed pdf
def embedding_pdf(filepath: str = "test.txt", index_name: str = index_name):
    document_store = load_store(index_name=index_name)
    pipeline = Pipeline()
    pipeline.add_component("converter",  PyPDFToDocument())
    pipeline.add_component("cleaner", DocumentCleaner())
    pipeline.add_component("splitter", DocumentSplitter(split_by=split_by, split_length=split_length, split_overlap = split_overlap))
    pipeline.add_component("embedder", SentenceTransformersDocumentEmbedder( model=model_name))
    # pipeline.add_component("embedder", CohereDocumentEmbedder (model=model_name))
    pipeline.add_component("writer", DocumentWriter(document_store=document_store, policy=DuplicatePolicy.OVERWRITE))
    pipeline.connect("converter", "cleaner")
    pipeline.connect("cleaner", "splitter")
    pipeline.connect("splitter","embedder")
    pipeline.connect("embedder", "writer")

    pipeline.run({"converter": {"sources": [filepath]}})

    return "Success!"



  
def embedding_csv(filepath: str = ".\courses.csv", index_name: str = index_name):
    df = pd.read_csv(filepath)
    size_col = len(df.columns)
    docu = []
    for index, row in df.iterrows():
        data_row= ""
        for i in range(size_col):
            data_row = f"{data_row}{df.columns[i]}:  {row.to_list()[i]}. "
        docu.append(Document(content=data_row))
    # init embedder
    # init qdrant cloud instance
    document_store = load_store(index_name=index_name)
    ## Use embedder Embedding file document for Fetch và Indexing
    # embedder = CohereDocumentEmbedder(model=model_name)
    embedder =  SentenceTransformersDocumentEmbedder( model=model_name)  
    embedder.warm_up()
    docs_with_embeddings =  embedder.run(docu)
    document_store.write_documents(docs_with_embeddings["documents"], policy=DuplicatePolicy.OVERWRITE)
    return "Success!"
    

def embedding_excel(file_path: str, index_name: str = index_name): 
# Looping through each file
  # Reading multiple sheets from an Excel file
  sheets_dict = pd.read_excel(file_path, engine="openpyxl", sheet_name=None)

  # Accessing individual sheets and displaying their contents
  for sheet_name, df in sheets_dict.items():
      print(f"Sheet '{sheet_name}':\n{df}\n")
      df.to_csv(f'{sheet_name}.csv')
      embedding_csv(f'{sheet_name}.csv',index_name)
  return "Success!"